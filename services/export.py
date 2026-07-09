"""Export service (design Section 5.7, TDD-14). Owner: Ankur2.

Renders an SOP into many formats. Machine-readable formats (JSON/XML/BPMN/testcases/rpa) are
generated from the process structure; document formats (md/html/docx/pdf) are human-facing.
"""
from __future__ import annotations

import base64
import io
import json
from collections.abc import Callable
from pathlib import Path
from xml.sax.saxutils import escape

from processiq_shared.models import SOP, SopStep

SUPPORTED = {"md", "markdown", "html", "json", "xml", "bpmn", "testcases", "rpa", "docx", "pdf"}

# image_loader(artifact_id) -> raw image bytes (or None). Provided by the API so document exports
# can embed each step's screenshot with the click-target box drawn on it.
ImageLoader = Callable[[str], "bytes | None"]


def export(sop: SOP, fmt: str, image_loader: ImageLoader | None = None) -> tuple[bytes, str]:
    """Return (bytes, content_type). Document formats embed annotated screenshots when a loader
    is supplied; without one they render text-only (used by tests)."""
    fmt = fmt.lower()
    if fmt in {"md", "markdown"}:
        return _markdown(sop, image_loader).encode(), "text/markdown"
    if fmt == "html":
        return _html(sop, image_loader).encode(), "text/html"
    if fmt == "json":
        return sop.model_dump_json(indent=2).encode(), "application/json"
    if fmt == "xml":
        return _xml(sop).encode(), "application/xml"
    if fmt == "bpmn":
        return _bpmn(sop).encode(), "application/xml"
    if fmt == "testcases":
        return _testcases(sop).encode(), "application/json"
    if fmt == "rpa":
        return _rpa(sop).encode(), "text/x-python"
    if fmt == "docx":
        return _docx(sop, image_loader), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf":
        return _pdf(sop, image_loader), "application/pdf"
    raise ValueError(f"unsupported format: {fmt}")


# ---------- annotated screenshots ----------
def _annotated_png(image_bytes: bytes, bbox: list[float] | None, max_w: int = 1000) -> bytes | None:
    """Draw the click-target box on the screenshot; downscale for embedding. None on failure."""
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None
    try:
        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    except Exception:
        return None
    if img.width > max_w:
        img = img.resize((max_w, round(img.height * max_w / img.width)))
    W, H = img.size
    if bbox and len(bbox) == 4:
        x, y, w, h = bbox
        if w > 0 and h > 0 and (w * h) < 0.9:  # skip empty sentinel / full-frame boxes
            box = [x * W, y * H, (x + w) * W, (y + h) * H]
            draw = ImageDraw.Draw(img)
            for off in range(3):  # thick, high-contrast rectangle
                draw.rectangle([box[0] - off, box[1] - off, box[2] + off, box[3] + off],
                               outline=(233, 69, 96))
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def _step_image(step: SopStep, image_loader: ImageLoader | None) -> bytes | None:
    if image_loader is None or step.screenshot_ref is None:
        return None
    raw = image_loader(step.screenshot_ref.artifact_id)
    if not raw:
        return None
    return _annotated_png(raw, step.screenshot_ref.bbox)


# ---------- shared document helpers (the required 7-section SOP layout) ----------
DOC_TITLE = "STANDARD OPERATING PROCEDURE (SOP)"


def _conf_pct(sop: SOP) -> int:
    return round(sop.overall_confidence * 100)


def _conf_label(pct: int) -> str:
    return "High" if pct >= 80 else "Moderate" if pct >= 60 else "Low"


def _conf_sentence(sop: SOP) -> str:
    pct = _conf_pct(sop)
    return (f"{pct}% — {_conf_label(pct)} confidence in the reconstructed workflow "
            f"(across {len(sop.steps)} step(s)).")


def _meta_block(sop: SOP) -> str:
    return (f"Version: {sop.version} | State: {sop.state.value} | "
            f"Models: {', '.join(sop.provenance.models) or 'n/a'}")


def _markdown(sop: SOP, image_loader: ImageLoader | None = None) -> str:
    L = [f"# {DOC_TITLE}", "", f"**Process:** {sop.title}", "", f"_{_meta_block(sop)}_", ""]
    L += [f"## 1. OBJECTIVE\n\n{sop.objective or '_Not specified._'}", ""]

    L += ["## 2. PRE-REQUISITES", ""]
    L += [f"- {p}" for p in sop.prerequisites] or ["- None"]

    L += ["", "## 3. STEP-BY-STEP PROCEDURE"]
    for s in sop.steps:  # each step: its description paragraph, then its screenshot, then the next
        flag = " ⚠️ needs review" if s.flags else ""
        L += ["", f"### Step {s.no}: {s.action}{flag}", "", s.description or ""]
        img = _step_image(s, image_loader)
        if img:
            b64 = base64.b64encode(img).decode()
            L += ["", f"![Step {s.no}](data:image/png;base64,{b64})"]

    L += ["", "## 4. EXCEPTION HANDLING", ""]
    L += [f"- {e}" for e in sop.exceptions] or ["- None documented."]

    L += ["", "## 5. VALIDATION & CHECKS", ""]
    L += [f"- {v}" for v in sop.validation] or ["- None documented."]

    L += ["", f"## 6. OUTPUT\n\n{sop.output or '_Not specified._'}", ""]
    L += ["## 7. CONFIDENCE SCORE", "", f"**{_conf_sentence(sop)}**"]
    return "\n".join(L)


def _html(sop: SOP, image_loader: ImageLoader | None = None) -> str:
    steps_html = []
    for s in sop.steps:  # each step: heading + description paragraph, then its screenshot
        badge = (" <span style='background:#e0a92b;color:#000;padding:0 6px;border-radius:9px;"
                 "font-size:11px'>needs review</span>") if s.flags else ""
        img = _step_image(s, image_loader)
        img_html = ""
        if img:
            b64 = base64.b64encode(img).decode()
            img_html = (f"<img src='data:image/png;base64,{b64}' "
                        f"style='max-width:100%;border:1px solid #ddd;border-radius:8px;margin-top:8px'/>")
        steps_html.append(
            f"<div style='margin:0 0 26px'>"
            f"<h3 style='margin:0 0 6px'>Step {s.no}: {escape(s.action)}{badge}</h3>"
            f"<p style='margin:0'>{escape(s.description).replace(chr(10), '<br>')}</p>"
            f"{img_html}</div>")

    def ul(items: list[str], empty: str) -> str:
        return "<ul>" + ("".join(f"<li>{escape(x)}</li>" for x in items) or f"<li>{empty}</li>") + "</ul>"

    return (f"<!doctype html><html><head><meta charset='utf-8'><title>{escape(sop.title)}</title>"
            f"<style>body{{font-family:Segoe UI,system-ui,sans-serif;max-width:860px;margin:24px auto;"
            f"padding:0 18px;color:#111}}h1{{margin-bottom:2px;font-size:22px}}h2{{margin-top:28px;"
            f"border-bottom:1px solid #eee;padding-bottom:4px}}h3{{margin-bottom:4px}}"
            f"</style></head><body>"
            f"<h1>{DOC_TITLE}</h1><p style='font-size:16px'><b>Process:</b> {escape(sop.title)}</p>"
            f"<p style='color:#6b7280;font-size:13px'><em>{escape(_meta_block(sop))}</em></p>"
            f"<h2>1. Objective</h2><p>{escape(sop.objective) or '<em>Not specified.</em>'}</p>"
            f"<h2>2. Pre-requisites</h2>{ul(sop.prerequisites, 'None')}"
            f"<h2>3. Step-by-Step Procedure</h2>{''.join(steps_html)}"
            f"<h2>4. Exception Handling</h2>{ul(sop.exceptions, 'None documented.')}"
            f"<h2>5. Validation &amp; Checks</h2>{ul(sop.validation, 'None documented.')}"
            f"<h2>6. Output</h2><p>{escape(sop.output) or '<em>Not specified.</em>'}</p>"
            f"<h2>7. Confidence Score</h2><p><b>{escape(_conf_sentence(sop))}</b></p>"
            f"</body></html>")


def _xml(sop: SOP) -> str:
    steps = "".join(
        f"<step no='{s.no}' confidence='{s.confidence}'><action>{escape(s.action)}</action>"
        f"<description>{escape(s.description)}</description></step>" for s in sop.steps
    )
    return (f"<?xml version='1.0' encoding='UTF-8'?><sop title='{escape(sop.title)}' "
            f"version='{sop.version}' confidence='{sop.overall_confidence}'>"
            f"<objective>{escape(sop.objective)}</objective><steps>{steps}</steps></sop>")


def _bpmn(sop: SOP) -> str:
    """Minimal but valid-shaped BPMN 2.0 with a task per step in sequence."""
    ns = "http://www.omg.org/spec/BPMN/20100524/MODEL"
    tasks, flows = [], []
    prev = "StartEvent_1"
    for s in sop.steps:
        tid = f"Task_{s.no}"
        tasks.append(f"<task id='{tid}' name='{escape(s.action)}'/>")
        flows.append(f"<sequenceFlow id='Flow_{s.no}' sourceRef='{prev}' targetRef='{tid}'/>")
        prev = tid
    flows.append(f"<sequenceFlow id='Flow_end' sourceRef='{prev}' targetRef='EndEvent_1'/>")
    return (f"<?xml version='1.0' encoding='UTF-8'?>"
            f"<definitions xmlns='{ns}' id='Defs_{sop.id}' targetNamespace='http://processiq'>"
            f"<process id='Process_{sop.id}' name='{escape(sop.title)}' isExecutable='false'>"
            f"<startEvent id='StartEvent_1'/>{''.join(tasks)}"
            f"<endEvent id='EndEvent_1'/>{''.join(flows)}</process></definitions>")


def _testcases(sop: SOP) -> str:
    cases = [{
        "id": f"TC-{s.no:03d}", "title": f"Verify: {s.action}",
        "steps": [s.description],
        "expected": "Step completes without error and UI advances as described.",
        "priority": "P1" if not s.flags else "P0",
    } for s in sop.steps]
    return json.dumps({"process": sop.title, "cases": cases}, indent=2)


def _rpa(sop: SOP) -> str:
    """RPA automation-script SKELETON (artifact only — never auto-executed, design §14.6)."""
    body = [f"# Auto-generated RPA skeleton for: {sop.title}",
            "# NOTE: artifact only. Review before running in an RPA runtime.",
            "def run(bot):"]
    for s in sop.steps:
        ref = s.screenshot_ref.artifact_id if s.screenshot_ref else "n/a"
        body.append(f"    bot.step({s.no}, action={s.action!r})  # ref={ref}")
    body.append("    return 'completed'")
    return "\n".join(body)


def _docx(sop: SOP, image_loader: ImageLoader | None = None) -> bytes:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(DOC_TITLE, level=0)
    doc.add_heading(f"Process: {sop.title}", level=1)
    doc.add_paragraph(_meta_block(sop)).italic = True

    doc.add_heading("1. Objective", level=1)
    doc.add_paragraph(sop.objective or "Not specified.")

    doc.add_heading("2. Pre-requisites", level=1)
    for p in sop.prerequisites or ["None"]:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("3. Step-by-Step Procedure", level=1)
    for s in sop.steps:  # each step: heading + description paragraph, then its screenshot
        flag = "  (needs review)" if s.flags else ""
        doc.add_heading(f"Step {s.no}: {s.action}{flag}", level=2)
        doc.add_paragraph(s.description or "")
        img = _step_image(s, image_loader)
        if img:
            doc.add_picture(io.BytesIO(img), width=Inches(6.0))

    doc.add_heading("4. Exception Handling", level=1)
    for e in sop.exceptions or ["None documented."]:
        doc.add_paragraph(e, style="List Bullet")

    doc.add_heading("5. Validation & Checks", level=1)
    for v in sop.validation or ["None documented."]:
        doc.add_paragraph(v, style="List Bullet")

    doc.add_heading("6. Output", level=1)
    doc.add_paragraph(sop.output or "Not specified.")

    doc.add_heading("7. Confidence Score", level=1)
    doc.add_paragraph(_conf_sentence(sop)).bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- branded PDF template (cover + header/footer chrome), owner: Pushp ----------
# Brand palette (matches the web UI / template.pdf).
_INK = "#0F172A"
_VIOLET = "#7C3AED"
_BLUE = "#2563EB"
_CYAN = "#06B6D4"
_SLATE = "#475569"
_MUTED = "#94A3B8"
_CARD_BG = "#F8FAFC"
_BORDER = "#E2E8F0"
_BADGE = "#EDE9FE"
_LOGO_PATH = Path(__file__).resolve().parents[1] / "apps" / "api" / "static" / "logo.jpeg"


def _trunc(text: str, n: int) -> str:
    text = " ".join((text or "").split())
    return text if len(text) <= n else text[: n - 1] + "…"


def _pdf(sop: SOP, image_loader: ImageLoader | None = None) -> bytes:
    from datetime import UTC, datetime

    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        Image,
        KeepTogether,
        ListFlowable,
        ListItem,
        NextPageTemplate,
        PageBreak,
        PageTemplate,
        Paragraph,
        Spacer,
    )

    W, H = landscape(A4)                 # 841.89 x 595.28
    INK, VIOLET, BLUE, CYAN = HexColor(_INK), HexColor(_VIOLET), HexColor(_BLUE), HexColor(_CYAN)
    SLATE, MUTED = HexColor(_SLATE), HexColor(_MUTED)
    CARD_BG, BORDER, BADGE = HexColor(_CARD_BG), HexColor(_BORDER), HexColor(_BADGE)
    now = datetime.now(UTC)
    pct = _conf_pct(sop)

    M = {
        "project": sop.title or "Untitled Process",
        "process_id": sop.process_id,
        "version": str(sop.version),
        "date_time": now.strftime("%d %b %Y, %H:%M UTC"),
        "date": now.strftime("%d %b %Y"),
        "status": sop.state.value,
        "confidence": f"{pct}% — {_conf_label(pct)}",
        "doc_id": sop.id,
        "client": _trunc((sop.tenant_id or "Internal").replace("_", " ").title(), 26),
    }

    def _logo(c, x, y, w):
        try:
            ir = ImageReader(str(_LOGO_PATH))
            iw, ih = ir.getSize()
            c.drawImage(ir, x, y, width=w, height=w * ih / iw, mask="auto")
        except Exception:  # noqa: BLE001 - logo is decorative; never fail the export
            pass

    def _footer(c):
        c.setFillColor(BLUE)
        c.rect(0, 0, W, 30, stroke=0, fill=1)
        c.setFillColor(HexColor("#DBEAFE"))
        c.setFont("Helvetica-Bold", 8)
        c.drawString(40, 11, "FROM VISUALS TO VALUE.  AUTOMATED.  INTELLIGENT.  ACTIONABLE.")
        c.drawRightString(W - 40, 11, "Confidential | Internal Use Only")

    def _circles(c, cx, cy):
        c.setStrokeColor(HexColor("#EEF2FF"))
        c.setLineWidth(1.2)
        for r in (34, 60, 88, 118):
            c.circle(cx, cy, r, stroke=1, fill=0)

    def draw_cover(c, _doc):
        # corner motif (mostly off-page → shows as a soft arc)
        c.setFillColor(VIOLET)
        c.circle(W + 20, -10, 90, stroke=0, fill=1)
        c.setFillColor(CYAN)
        c.circle(W - 8, -22, 55, stroke=0, fill=1)
        _circles(c, W - 18, H - 26)
        _footer(c)
        _logo(c, 40, 486, 238)
        c.setStrokeColor(CYAN)
        c.setLineWidth(3)
        c.line(40, 470, 430, 470)
        # title: "SOP DOCUMENT"
        c.setFont("Helvetica-Bold", 38)
        c.setFillColor(INK)
        c.drawString(40, 412, "SOP ")
        c.setFillColor(VIOLET)
        c.drawString(40 + c.stringWidth("SOP ", "Helvetica-Bold", 38), 412, "DOCUMENT")
        c.setFont("Helvetica", 13)
        c.setFillColor(SLATE)
        c.drawString(40, 390, "AI-generated Standard Operating Procedure")
        c.setStrokeColor(VIOLET)
        c.setLineWidth(2.5)
        c.line(40, 381, 320, 381)
        c.setFillColor(MUTED)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(40, 349, "P R O J E C T   T I T L E")
        c.setFillColor(INK)
        c.setFont("Helvetica-Bold", 24)
        c.drawString(40, 317, _trunc(M["project"], 34))
        # metadata card
        cx, cy, cw, ch = 40, 58, 430, 232
        c.setFillColor(CARD_BG)
        c.setStrokeColor(BORDER)
        c.roundRect(cx, cy, cw, ch, 12, stroke=1, fill=1)
        rows = [("Process ID", M["process_id"]), ("Version", M["version"]),
                ("Generated On", M["date_time"]), ("Generated By", "ProcessIQ AI Engine"),
                ("Review Status", M["status"]), ("Overall Confidence", M["confidence"]),
                ("Document ID", M["doc_id"])]
        for i, (label, val) in enumerate(rows):
            ry = cy + ch - 26 - i * 27
            c.setFillColor(VIOLET)
            c.circle(cx + 22, ry + 3, 2.4, stroke=0, fill=1)
            c.setFillColor(INK)
            c.setFont("Helvetica-Bold", 9.5)
            c.drawString(cx + 34, ry, label)
            c.setFillColor(SLATE)
            c.setFont("Helvetica", 9.5)
            c.drawString(cx + 158, ry, ": " + _trunc(str(val), 32))
        # right "prepared for / by" cards
        def prep_card(y, badge, eyebrow, value):
            c.setFillColor(HexColor("#FFFFFF"))
            c.setStrokeColor(BORDER)
            c.roundRect(505, y, 297, 100, 12, stroke=1, fill=1)
            c.setFillColor(BADGE)
            c.roundRect(520, y + 25, 50, 50, 10, stroke=0, fill=1)
            c.setFillColor(VIOLET)
            c.setFont("Helvetica-Bold", 15)
            c.drawCentredString(545, y + 43, badge)
            c.setFillColor(MUTED)
            c.setFont("Helvetica-Bold", 8)
            c.drawString(585, y + 60, eyebrow)
            c.setFillColor(INK)
            c.setFont("Helvetica-Bold", 13)
            c.drawString(585, y + 40, value)
        prep_card(185, "TO", "PREPARED FOR", M["client"])
        prep_card(70, "AI", "PREPARED BY", "ProcessIQ AI")
        c.setFillColor(MUTED)
        c.setFont("Helvetica", 8)
        c.drawString(40, 42, "Confidential  |  Internal Use Only")

    def draw_content(c, _doc):
        _footer(c)
        _circles(c, W - 18, H - 26)
        _logo(c, 40, H - 52, 118)
        cols = [("Process Name", _trunc(M["project"], 22)), ("Process ID", _trunc(M["process_id"], 16)),
                ("Version", M["version"]), ("Date", M["date"]), ("Classification", "Internal")]
        xs = [195, 360, 470, 545, 640]
        for x, (label, val) in zip(xs, cols):
            c.setFillColor(VIOLET)
            c.setFont("Helvetica-Bold", 7.5)
            c.drawString(x, H - 32, label.upper())
            c.setFillColor(INK)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(x, H - 46, val)
        c.setStrokeColor(BORDER)
        c.setLineWidth(1)
        c.line(40, H - 62, W - 40, H - 62)
        # rounded content container
        c.setStrokeColor(BORDER)
        c.roundRect(30, 40, W - 60, H - 62 - 44, 14, stroke=1, fill=0)

    # page-number badge (drawn once the total is known)
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *a, **k):
            super().__init__(*a, **k)
            self._pages: list[dict] = []

        def showPage(self):
            self._pages.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._pages)
            for state in self._pages:
                self.__dict__.update(state)
                if self._pageNumber > 1:      # cover has no page badge
                    self.setFillColor(VIOLET)
                    self.roundRect(W - 118, H - 52, 78, 32, 7, stroke=0, fill=1)
                    self.setFillColor(HexColor("#FFFFFF"))
                    self.setFont("Helvetica", 7)
                    self.drawCentredString(W - 79, H - 30, "PAGE")
                    self.setFont("Helvetica-Bold", 10)
                    self.drawCentredString(W - 79, H - 43, f"{self._pageNumber} of {total}")
                super().showPage()
            super().save()

    # ---- styles + flowables (the 7-section SOP body) ----
    h_sec = ParagraphStyle("Sec", fontName="Helvetica-Bold", fontSize=13, textColor=VIOLET,
                           spaceBefore=12, spaceAfter=5, leading=16)
    h_step = ParagraphStyle("Step", fontName="Helvetica-Bold", fontSize=11, textColor=INK,
                            spaceBefore=8, spaceAfter=2, leading=14)
    body = ParagraphStyle("Body", fontName="Helvetica", fontSize=10, textColor=HexColor("#1F2937"),
                          leading=14)
    flag_st = ParagraphStyle("Flag", parent=h_step, textColor=HexColor("#B45309"))

    frame_w = W - 96
    frame_h = (H - 90) - 58

    def bullets(items: list[str], empty: str) -> ListFlowable:
        items = items or [empty]
        return ListFlowable([ListItem(Paragraph(escape(x), body)) for x in items],
                            bulletType="bullet", leftIndent=14)

    flow = [NextPageTemplate("content"), PageBreak(),
            Paragraph("1. Objective", h_sec),
            Paragraph(escape(sop.objective) or "<i>Not specified.</i>", body), Spacer(1, 4),
            Paragraph("2. Pre-requisites", h_sec), bullets(sop.prerequisites, "None"),
            Paragraph("3. Step-by-Step Procedure", h_sec)]

    for s in sop.steps:  # each step: heading + description paragraph, then its screenshot
        style = flag_st if s.flags else h_step
        flag = "  (needs review)" if s.flags else ""
        block = [Paragraph(f"Step {s.no}: {escape(s.action)}{flag}", style),
                 Paragraph(escape(s.description).replace("\n", "<br/>"), body)]
        img = _step_image(s, image_loader)
        if img:
            iw, ih = ImageReader(io.BytesIO(img)).getSize()
            w = min(frame_w, iw)
            h = ih * w / iw
            if h > frame_h - 78:              # leave room for the heading + description above it
                h = frame_h - 78
                w = iw * h / ih
            block += [Spacer(1, 4), Image(io.BytesIO(img), width=w, height=h)]
        # keep a step's heading, description and screenshot together on one page
        flow += [KeepTogether(block), Spacer(1, 8)]

    flow += [Paragraph("4. Exception Handling", h_sec), bullets(sop.exceptions, "None documented."),
             Paragraph("5. Validation & Checks", h_sec), bullets(sop.validation, "None documented."),
             Paragraph("6. Output", h_sec),
             Paragraph(escape(sop.output) or "<i>Not specified.</i>", body),
             Paragraph("7. Confidence Score", h_sec),
             Paragraph(f"<b>{escape(_conf_sentence(sop))}</b>", body)]

    buf = io.BytesIO()
    doc = BaseDocTemplate(buf, pagesize=(W, H), leftMargin=48, rightMargin=48,
                          topMargin=90, bottomMargin=58)
    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[Frame(40, 40, W - 80, H - 80, id="cover")],
                     onPage=draw_cover),
        PageTemplate(id="content", frames=[Frame(48, 58, frame_w, frame_h, id="content")],
                     onPage=draw_content),
    ])
    doc.build(flow, canvasmaker=NumberedCanvas)
    return buf.getvalue()
